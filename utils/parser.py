"""
parser.py
---------
Resume parsing utilities for ResumeFit AI.
Extracts structured data from PDF and DOCX resumes.
"""

import re
import io
import tempfile
import os
from typing import Optional

# PDF parsing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─────────────────────────────────────────────
# Text Extraction
# ─────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract raw text from a PDF file using pdfplumber (preferred) or PyPDF2 fallback.
    Returns the full text as a single string.
    """
    text = ""

    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception:
            pass  # Fall through to PyPDF2

    if PYPDF2_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            return f"Error extracting PDF text: {str(e)}"

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract raw text from a DOCX file using python-docx.
    Returns the full text as a single string.
    """
    if not DOCX_AVAILABLE:
        return "python-docx not available."

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"Error extracting DOCX text: {str(e)}"


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Route to the correct extractor based on file type.
    file_type should be 'pdf' or 'docx'.
    """
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        return ""


# ─────────────────────────────────────────────
# Field Extraction
# ─────────────────────────────────────────────

def extract_email(text: str) -> str:
    """Extract the first email address found in text."""
    pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract the first phone number found in text."""
    pattern = r"(\+?\d[\d\s\-().]{7,}\d)"
    match = re.search(pattern, text)
    return match.group(0).strip() if match else ""


def extract_name(text: str) -> str:
    """
    Heuristic: the name is usually on the first non-empty line
    that contains only letters and spaces (no @, digits, etc.).
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:6]:  # Check top 6 lines
        # A name is typically 2-4 words, all letters
        words = line.split()
        if 1 < len(words) <= 5 and all(re.match(r"^[A-Za-z.\-']+$", w) for w in words):
            return line
    return lines[0] if lines else ""


def extract_linkedin(text: str) -> str:
    """Extract LinkedIn profile URL if present."""
    pattern = r"linkedin\.com/in/[A-Za-z0-9\-_%]+"
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else ""


def extract_github(text: str) -> str:
    """Extract GitHub profile URL if present."""
    pattern = r"github\.com/[A-Za-z0-9\-_%]+"
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else ""


def extract_experience_years(text: str) -> int:
    """
    Estimate years of experience by searching for explicit mentions
    like '5 years of experience' or '3+ years'.
    """
    patterns = [
        r"(\d+)\+?\s*years?\s+of\s+(?:professional\s+)?experience",
        r"(\d+)\+?\s*years?\s+experience",
        r"experience\s+of\s+(\d+)\+?\s*years?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    return 0


def extract_education(text: str) -> list[str]:
    """
    Extract education entries based on degree keywords.
    Returns a list of matching lines.
    """
    degree_keywords = [
        "bachelor", "master", "phd", "ph.d", "doctorate", "b.sc", "m.sc",
        "b.tech", "m.tech", "mba", "bca", "mca", "be", "me", "bs", "ms",
        "associate", "diploma", "certificate", "high school", "matric", "intermediate"
    ]
    education_lines = []
    lines = text.splitlines()
    for line in lines:
        if any(kw in line.lower() for kw in degree_keywords):
            education_lines.append(line.strip())
    return list(dict.fromkeys(education_lines))  # Remove duplicates


def extract_certifications(text: str) -> list[str]:
    """
    Extract certification mentions from text.
    """
    cert_keywords = [
        "certified", "certification", "certificate", "aws certified",
        "google cloud", "azure", "pmp", "cissp", "ccna", "cpa", "cfa",
        "tensorflow", "coursera", "udemy", "microsoft certified"
    ]
    cert_lines = []
    lines = text.splitlines()
    for line in lines:
        if any(kw in line.lower() for kw in cert_keywords):
            cert_lines.append(line.strip())
    return list(dict.fromkeys(cert_lines))


def extract_section(text: str, section_names: list[str]) -> str:
    """
    Extract the content of a specific resume section.
    section_names: list of possible headings (e.g. ['experience', 'work history'])
    Returns text from the section heading until the next heading.
    """
    section_pattern = r"(?:^|\n)(" + "|".join(re.escape(s) for s in section_names) + r")\s*[:.\-]?\s*\n"
    next_section_pattern = r"\n[A-Z][A-Z\s]{2,30}(?:\s*[:.\-])?\s*\n"

    match = re.search(section_pattern, text, re.IGNORECASE)
    if not match:
        return ""

    start = match.end()
    # Find where the next section starts
    next_match = re.search(next_section_pattern, text[start:])
    end = start + next_match.start() if next_match else len(text)
    return text[start:end].strip()


# ─────────────────────────────────────────────
# Master Parser
# ─────────────────────────────────────────────

def parse_resume(file_bytes: bytes, file_type: str) -> dict:
    """
    Full resume parser. Extracts all key fields from resume bytes.
    Returns a dictionary with structured resume data.
    """
    raw_text = extract_text(file_bytes, file_type)

    if not raw_text or raw_text.startswith("Error"):
        return {
            "raw_text": raw_text,
            "name": "",
            "email": "",
            "phone": "",
            "linkedin": "",
            "github": "",
            "education": [],
            "certifications": [],
            "experience_years": 0,
            "parse_error": raw_text if raw_text.startswith("Error") else "Empty file",
        }

    return {
        "raw_text": raw_text,
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "linkedin": extract_linkedin(raw_text),
        "github": extract_github(raw_text),
        "education": extract_education(raw_text),
        "certifications": extract_certifications(raw_text),
        "experience_years": extract_experience_years(raw_text),
        "parse_error": None,
    }
